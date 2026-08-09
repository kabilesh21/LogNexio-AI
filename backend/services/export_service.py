import io
import json
from datetime import datetime
from typing import Dict, Any

# ReportLab imports for PDF export
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, KeepTogether, Image
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.graphics.shapes import Drawing, Circle, Line, Rect

# python-docx imports for Word export
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from models.ai_report import AIReport
from utils.analysis_logger import get_analysis_logger

logger = get_analysis_logger("ExportService")


class ExportService:
    """
    Generates professional PDF, DOCX, Markdown, and JSON export formats for AI incident reports.
    """

    @staticmethod
    def export_json(report: AIReport) -> str:
        """Exports report as formatted JSON string."""
        return json.dumps(report.dict(), indent=2, ensure_ascii=False)

    @staticmethod
    def export_markdown(report: AIReport) -> str:
        """Exports report as GitHub-flavored Markdown document."""
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        res_steps = "\n".join([f"{i+1}. {step}" for i, step in enumerate(report.resolution_steps)])
        prev_steps = "\n".join([f"- {m}" for m in report.preventive_measures])
        keywords_str = ", ".join([f"`{k}`" for k in report.keywords])

        md = f"""# 🛡️ LogNexio AI Incident Analysis Report

> **Incident ID:** `{report.incident_id}`  
> **Generated Date:** {now}  
> **Classification:** `{report.error_type}`  
> **Severity:** **{report.severity}**  
> **AI Confidence:** {report.confidence} | **Est. Fix Time:** {report.estimated_fix_time}

---

## 📋 Executive Summary
{report.incident_summary}

---

## 🔍 Root Cause
{report.root_cause}

---

## 💻 Technical Explanation
{report.technical_explanation}

---

## 🏢 Business Impact & System Risk
{report.business_impact}

---

## 🛠️ Actionable Resolution Steps
{res_steps}

---

## 🛡️ Recommended Preventive Measures
{prev_steps}

---

## 🏷️ Metadata & Keywords
- **Affected Component:** `{report.affected_component}`
- **Technical Keywords:** {keywords_str}

---
*Generated automatically by LogNexio AI Enterprise Platform.*
"""
        return md.strip()

    @staticmethod
    def export_docx(report: AIReport) -> bytes:
        """Generates a professional Microsoft Word (.docx) binary document."""
        doc = Document()

        # Page Setup
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        # Styles
        style_normal = doc.styles['Normal']
        style_normal.font.name = 'Calibri'
        style_normal.font.size = Pt(11)
        style_normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Header Title
        title_p = doc.add_paragraph()
        title_run = title_p.add_run("LOGNEXIO AI — INCIDENT ANALYSIS REPORT")
        title_run.bold = True
        title_run.font.size = Pt(20)
        title_run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

        subtitle_p = doc.add_paragraph()
        sub_run = subtitle_p.add_run(f"Incident ID: {report.incident_id}  |  Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        sub_run.font.size = Pt(9.5)
        sub_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        doc.add_paragraph() # Spacer

        # Metadata Summary Table
        table = doc.add_table(rows=4, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        meta_data = [
            ("Error Type", report.error_type),
            ("Severity", report.severity),
            ("Affected Component", report.affected_component),
            ("AI Confidence / Fix Time", f"{report.confidence} confidence  |  Est. Fix: {report.estimated_fix_time}"),
        ]

        for i, (label, val) in enumerate(meta_data):
            row_cells = table.rows[i].cells
            row_cells[0].text = label
            row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[0].paragraphs[0].runs[0].font.size = Pt(10)
            row_cells[0].width = Inches(2.0)
            
            row_cells[1].text = str(val)
            row_cells[1].paragraphs[0].runs[0].font.size = Pt(10)
            row_cells[1].width = Inches(4.5)

        doc.add_paragraph()

        def add_section_header(title_text):
            h = doc.add_heading(level=2)
            run = h.add_run(title_text)
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
            run.bold = True

        # Executive Summary
        add_section_header("1. Executive Summary")
        p = doc.add_paragraph(report.incident_summary)
        p.paragraph_format.space_after = Pt(12)

        # Root Cause
        add_section_header("2. Root Cause Analysis")
        p = doc.add_paragraph(report.root_cause)
        p.paragraph_format.space_after = Pt(12)

        # Technical Explanation
        add_section_header("3. Technical Explanation")
        p = doc.add_paragraph(report.technical_explanation)
        p.paragraph_format.space_after = Pt(12)

        # Business Impact
        add_section_header("4. Business Impact & Risk")
        p = doc.add_paragraph(report.business_impact)
        p.paragraph_format.space_after = Pt(12)

        # Resolution Steps
        add_section_header("5. Recommended Resolution Steps")
        for i, step in enumerate(report.resolution_steps, 1):
            p = doc.add_paragraph(f"{i}. {step}", style='List Number')
            p.paragraph_format.space_after = Pt(4)

        doc.add_paragraph()

        # Preventive Measures
        add_section_header("6. Preventive Measures")
        for measure in report.preventive_measures:
            p = doc.add_paragraph(measure, style='List Bullet')
            p.paragraph_format.space_after = Pt(4)

        # Footer
        section = doc.sections[0]
        footer = section.footer
        f_p = footer.paragraphs[0]
        f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f_run = f_p.add_run("Generated by LogNexio AI Enterprise Platform")
        f_run.font.size = Pt(8.5)
        f_run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

        # Save to BytesIO stream
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    @staticmethod
    def export_pdf(report: AIReport) -> bytes:
        """Generates a styled PDF report using ReportLab."""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=36,
            leftMargin=36,
            topMargin=36,
            bottomMargin=36
        )

        styles = getSampleStyleSheet()

        # Custom Palette (Sepia / Warm Light Brown Theme)
        BROWN_DARK = colors.HexColor("#1C120C")
        BROWN_PRIMARY = colors.HexColor("#8C5630")
        BROWN_MUTED = colors.HexColor("#70573E")
        BROWN_BG = colors.HexColor("#FAF6F0")
        BROWN_BORDER = colors.HexColor("#DFCBB5")

        # Custom Paragraph Styles
        title_style = ParagraphStyle(
            'DocTitle',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=16,
            leading=20,
            textColor=BROWN_DARK,
            spaceAfter=4
        )
        subtitle_style = ParagraphStyle(
            'DocSubTitle',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=9,
            leading=12,
            textColor=BROWN_MUTED,
            spaceAfter=12
        )
        h2_style = ParagraphStyle(
            'SectionHeader',
            parent=styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=11,
            leading=15,
            textColor=BROWN_PRIMARY,
            spaceBefore=12,
            spaceAfter=6
        )
        body_style = ParagraphStyle(
            'BodyDark',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=9,
            leading=13,
            textColor=BROWN_DARK,
            spaceAfter=6
        )
        table_cell = ParagraphStyle(
            'TableCell',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=8.5,
            leading=11.5,
            textColor=BROWN_DARK
        )
        table_cell_bold = ParagraphStyle(
            'TableCellBold',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=8.5,
            leading=11.5,
            textColor=BROWN_DARK
        )

        story = []

        # 1. Header Logo & 
        logo_text_style = ParagraphStyle(
            'LogoText',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=22,
            leading=26,
            textColor=BROWN_PRIMARY,
            alignment=1 # Center aligned
        )
        logo_sub_style = ParagraphStyle(
            'LogoSub',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=7.5,
            leading=10,
            textColor=BROWN_MUTED,
            spaceBefore=1,
            alignment=1 # Center aligned
        )
        addr_style = ParagraphStyle(
            'CompanyAddr',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=7.5,
            leading=11,
            textColor=BROWN_MUTED,
            alignment=1 # Center aligned
        )

        # Load Logo image from backend directory
        import os
        from PIL import Image as PILImage
        current_dir = os.path.dirname(os.path.abspath(__file__))
        logo_path = os.path.join(current_dir, "..", "logo_icon.png")
        flat_logo_path = os.path.join(current_dir, "..", "logo_icon_flat.png")
        
        logo_loaded = False
        if os.path.exists(logo_path):
            try:
                img = PILImage.open(logo_path)
                if img.mode in ('RGBA', 'LA') or (img.mode == 'P' and 'transparency' in img.info):
                    alpha = img.convert('RGBA')
                    # Pure white background to match the white page background
                    bg = PILImage.new('RGB', img.size, (255, 255, 255))
                    bg.paste(alpha, mask=alpha.split()[3])
                    try:
                        bg.save(flat_logo_path, "PNG")
                        logo_img_path = flat_logo_path
                    except Exception:
                        # Fallback for read-only filesystem (e.g. Vercel)
                        temp_flat_path = os.path.join("/tmp", "logo_icon_flat.png")
                        try:
                            # Create directory if it doesn't exist
                            os.makedirs("/tmp", exist_ok=True)
                        except Exception:
                            pass
                        bg.save(temp_flat_path, "PNG")
                        logo_img_path = temp_flat_path
                else:
                    logo_img_path = logo_path
                
                # Size is square (394x385). 1.15 inches is perfect!
                logo_img = Image(logo_img_path, width=1.15*inch, height=1.12*inch)
                logo_img.hAlign = 'CENTER'
                story.append(logo_img)
                story.append(Spacer(1, 4))
                logo_loaded = True
            except Exception as e:
                logger.error(f"Error processing logo icon: {e}")

        # Append centered brand name, sub title and address
        story.append(Paragraph("LOGNEXIO AI", logo_text_style))
        story.append(Paragraph("INTELLIGENCE THAT INSPIRES", logo_sub_style))
        story.append(Paragraph("LogNexio Technologies Inc. &nbsp;|&nbsp; contact@lognexio.ai &nbsp;|&nbsp; San Jose, CA", addr_style))
        story.append(Spacer(1, 6))

        story.append(HRFlowable(width="100%", thickness=1.5, color=BROWN_PRIMARY, spaceAfter=12))

        # Title
        story.append(Paragraph("INCIDENT ANALYSIS REPORT", title_style))
        story.append(Paragraph(f"Incident ID: <b>{report.incident_id}</b> | Generated Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}", subtitle_style))

        # Metadata Table
        meta_table_data = [
            [
                Paragraph("Error Type:", table_cell_bold),
                Paragraph(report.error_type, table_cell),
                Paragraph("Severity:", table_cell_bold),
                Paragraph(f"<b>{report.severity}</b>", table_cell),
            ],
            [
                Paragraph("Component:", table_cell_bold),
                Paragraph(report.affected_component, table_cell),
                Paragraph("Confidence:", table_cell_bold),
                Paragraph(f"{report.confidence} (Fix: {report.estimated_fix_time})", table_cell),
            ]
        ]
        meta_table = Table(meta_table_data, colWidths=[80, 180, 80, 180])
        meta_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), BROWN_BG),
            ('BOX', (0,0), (-1,-1), 0.5, BROWN_BORDER),
            ('INNERGRID', (0,0), (-1,-1), 0.5, BROWN_BORDER),
            ('TOPPADDING', (0,0), (-1,-1), 5),
            ('BOTTOMPADDING', (0,0), (-1,-1), 5),
            ('LEFTPADDING', (0,0), (-1,-1), 8),
            ('RIGHTPADDING', (0,0), (-1,-1), 8),
        ]))
        story.append(meta_table)
        story.append(Spacer(1, 8))

        # 1. Executive Summary
        story.append(Paragraph("1. Executive Summary", h2_style))
        story.append(Paragraph(report.incident_summary, body_style))

        # 2. Root Cause
        story.append(Paragraph("2. Root Cause", h2_style))
        story.append(Paragraph(report.root_cause, body_style))

        # 3. Technical Explanation
        story.append(Paragraph("3. Technical Explanation", h2_style))
        story.append(Paragraph(report.technical_explanation, body_style))

        # 4. Business Impact
        story.append(Paragraph("4. Business Impact & Risk", h2_style))
        story.append(Paragraph(report.business_impact, body_style))

        # 5. Resolution Steps
        story.append(Paragraph("5. Recommended Resolution Steps", h2_style))
        for i, step in enumerate(report.resolution_steps, 1):
            story.append(Paragraph(f"<b>{i}.</b> {step}", body_style))

        # 6. Preventive Measures
        story.append(Paragraph("6. Recommended Preventive Measures", h2_style))
        for measure in report.preventive_measures:
            story.append(Paragraph(f"• {measure}", body_style))

        story.append(Spacer(1, 10))
        story.append(HRFlowable(width="100%", thickness=0.5, color=BROWN_BORDER, spaceAfter=8))
        story.append(Paragraph("<font color='#70573E' size='7.5'>Generated by LogNexio AI Enterprise Incident Analytics Platform</font>", ParagraphStyle('Footer', alignment=1)))

        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
